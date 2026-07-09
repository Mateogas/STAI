from __future__ import annotations

import argparse
import json
import re
from copy import deepcopy
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "docs" / "TECHNICAL_WRITEUP.md"
TEMPLATE_DOCX = Path(r"C:\Users\Zurax\Downloads\[STAI] Technical Write-up.docx")
OUTPUT_DOCX = ROOT / "docs" / "STAI_AISHA_Technical_Writeup.docx"

BDO_BLUE = "003B7A"
BDO_DARK_BLUE = "002D62"
BDO_GOLD = "FDB913"
LIGHT_BLUE = "EAF2FB"
LIGHT_GOLD = "FFF4CF"
LIGHT_GRAY = "F5F7FA"
MID_GRAY = "D9E2EC"
INK = "1F2933"
MUTED = "5B6770"
WHITE = "FFFFFF"

CONTENT_WIDTH = 9360


def rgb(hex_color: str) -> RGBColor:
    h = hex_color.strip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def remove_element(element):
    element.getparent().remove(element)


def set_run_font(run, name="Arial", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color) if isinstance(color, str) else color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold=False, color=INK, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0, line=1.1)
    if align is not None:
        p.alignment = align
    for idx, part in enumerate(text.split("\n")):
        if idx:
            p.add_run().add_break()
        run = p.add_run(part)
        set_run_font(run, size=size, color=color, bold=bold)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=MID_GRAY, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def apply_table_style(table):
    try:
        table.style = "Table Grid"
    except KeyError:
        pass


def set_table_geometry(table, widths):
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    old_grid = tbl.tblGrid
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(0, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    hdr = OxmlElement("w:tblHeader")
    hdr.set(qn("w:val"), "true")
    tr_pr.append(hdr)


def add_field(paragraph, instruction: str, display_text: str = ""):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    if display_text:
        text = OxmlElement("w:t")
        text.text = display_text
        run._r.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
    return run


def add_bookmark(paragraph, name: str, bookmark_id: int):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_hyperlink(paragraph, text: str, anchor: str, color=BDO_BLUE):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    r_pr.append(color_el)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.8)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BDO_BLUE, 14, 8),
        ("Heading 2", 13, BDO_BLUE, 10, 5),
        ("Heading 3", 11.5, BDO_DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Caption" in styles:
        cap = styles["Caption"]
        cap.font.name = "Arial"
        cap._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        cap._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        cap.font.size = Pt(9)
        cap.font.italic = True
        cap.font.color.rgb = rgb(MUTED)
        cap.paragraph_format.space_before = Pt(3)
        cap.paragraph_format.space_after = Pt(8)


def setup_document_settings(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    section.different_first_page_header_footer = True

    footer = section.footer.paragraphs[0]
    clear_paragraph(footer)
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("AISHA Technical Write-Up | Page ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(footer, "PAGE", "1")

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def keep_cover_and_clear_body(doc):
    body = doc._body._element
    keep_until = doc.paragraphs[21]._p
    found = False
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        if child is keep_until:
            found = True
            continue
        if found:
            remove_element(child)


def update_cover(doc):
    replacements = {
        1: "De La Salle University-Manila\nCollege of Computer Studies",
        2: "2401 Taft Avenue Malate, Manila",
        3: "Term X, A.Y. 202X - 202X",
        4: "\nAISHA: AI Support for Hires and Associates",
        5: "Technical Write-Up",
        7: "In partial fulfillment of the course requirements for",
        8: "<COURSE> <SECTION> - <FULL COURSE CODE>",
        10: "Submitted by:",
        11: "<Member 1>",
        12: "<GROUP MEMBERS>",
        19: "Submitted to:",
        20: "<Prof_Name>, PhD",
    }
    for idx, text in replacements.items():
        p = doc.paragraphs[idx]
        style = p.style
        alignment = p.alignment
        clear_paragraph(p)
        for j, part in enumerate(text.split("\n")):
            if j:
                p.add_run().add_break()
            run = p.add_run(part)
            set_run_font(run, name="Arial")
        p.style = style
        p.alignment = alignment

    # Apply restrained title color on the cover only.
    for idx in (4, 5):
        for run in doc.paragraphs[idx].runs:
            set_run_font(run, color=BDO_BLUE, bold=True)
    doc.paragraphs[5].runs[0].font.size = Pt(16)

    p = doc.paragraphs[21]
    clear_paragraph(p)
    r = p.add_run(date.today().strftime("%B %d, %Y"))
    set_run_font(r, name="Arial")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_break(WD_BREAK.PAGE)


def add_para(doc, text="", *, style=None, align=None, before=0, after=6, bold=False, italic=False, color=INK, size=10.8):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    set_paragraph_spacing(p, before=before, after=after)
    add_inline_runs(p, text, bold=bold, italic=italic, color=color, size=size)
    return p


INLINE_RE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")


def add_inline_runs(paragraph, text: str, *, bold=False, italic=False, color=INK, size=10.8):
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=size, color=color, bold=bold, italic=italic)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=max(size - 1, 8.5), color=BDO_DARK_BLUE)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True, italic=italic)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, color=color, bold=bold, italic=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def add_caption(doc, label: str, caption: str):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=3, after=8, line=1.1)
    run = p.add_run(f"{label}. {caption}")
    set_run_font(run, size=9, color=MUTED, italic=True)
    return p


def add_contents(doc, page_map: dict[str, str] | None):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("Table of Contents")
    add_bookmark(p, "toc", 1)
    intro = add_para(
        doc,
        "This contents page is intentionally simple so it remains easy to edit in Word. Page numbers are filled from the rendered draft where available.",
        size=9.5,
        color=MUTED,
        after=8,
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT

    rows = [
        ("1. Business Case", "Why AISHA matters for Alyssa's time-to-ramp"),
        ("2. Methodology", "How the local-first prototype was built"),
        ("3. Architecture", "Inputs, process, outputs, tools, safety, and monitoring"),
        ("4. Experiments and Evaluation", "What was tested, what passed, what failed, and what we learned"),
        ("5. Retrospective", "Human reflection on what worked and what we would change"),
        ("References", "Sources cited in the write-up"),
        ("Appendix A. Implementation Evidence Map", "Supporting implementation and test evidence"),
    ]
    table = doc.add_table(rows=1, cols=3)
    apply_table_style(table)
    set_table_geometry(table, [3820, 4380, 1160])
    set_table_borders(table)
    headers = ["Section", "Focus", "Page"]
    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, BDO_BLUE)
        set_cell_text(cell, h, bold=True, color=WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    repeat_table_header(table.rows[0])
    for title, focus in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], title, bold=True, color=BDO_DARK_BLUE, size=9.2)
        set_cell_text(cells[1], focus, color=INK, size=9.2)
        set_cell_text(cells[2], str((page_map or {}).get(title, "")), color=INK, size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER)
        for c in cells:
            set_cell_shading(c, WHITE)
    add_page_break(doc)


def add_markdown_table(doc, rows: list[list[str]], caption_label: str | None = None):
    if not rows:
        return
    cols = len(rows[0])
    if caption_label:
        add_caption(doc, caption_label, "")
    table = doc.add_table(rows=1, cols=cols)
    apply_table_style(table)
    if cols == 5:
        widths = [1500, 1900, 1900, 1900, 2160]
    elif cols == 3:
        widths = [2200, 3560, 3600]
    else:
        widths = [CONTENT_WIDTH // cols] * cols
    set_table_geometry(table, widths)
    set_table_borders(table)
    for j, text in enumerate(rows[0]):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, BDO_BLUE)
        set_cell_text(cell, text, bold=True, color=WHITE, size=8.7, align=WD_ALIGN_PARAGRAPH.CENTER)
    repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows[1:]):
        cells = table.add_row().cells
        for j, text in enumerate(row):
            set_cell_shading(cells[j], LIGHT_BLUE if r_idx % 2 == 0 else WHITE)
            align = WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[j], text, color=INK, size=8.3, align=align)
    doc.add_paragraph()


def add_figure1(doc):
    add_caption(doc, "Figure 1", "Input-process-output system architecture")
    table = doc.add_table(rows=1, cols=3)
    apply_table_style(table)
    set_table_geometry(table, [3000, 3360, 3000])
    set_table_borders(table, color="C7D4E2", size="8")
    headers = ["INPUT", "PROCESS", "OUTPUT"]
    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, BDO_BLUE)
        set_cell_text(cell, h, bold=True, color=WHITE, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        (
            "Static Knowledge\n- docs .md\n- policy docs",
            "1. Knowledge Preparation\n- ingest\n- chunk\n- embed\n- Chroma vector store",
            "User-Facing Output\n- grounded response\n- citations\n- ramp plan\n- task updates",
        ),
        (
            "Seeded Data\n- employee JSON\n- plan JSON\n- org JSON",
            "2. Context Assembly\n- retrieve knowledge\n- load user state\n- load memory/history",
            "HR/Admin Output\n- escalations\n- pulse trends\n- HR dashboard",
        ),
        (
            "Persistent State\n- chat memory\n- pulse history\n- escalations",
            "3. Agent Reasoning + Tools\n- guardrail classifier\n- LangGraph/ReAct agent\n- local LLM\n- agent tools",
            "System Side Effects\n- updated memory\n- updated SQLite\n- JSONL logs",
        ),
        (
            "Runtime Input\n- user message\n- demo clock",
            "4. Safety + Monitoring\n- output guardrails\n- PII redaction\n- observability",
            "Supported Loop\n- new-hire support\n- HR signal\n- local audit trail",
        ),
    ]
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for j, text in enumerate(row):
            set_cell_shading(cells[j], LIGHT_BLUE if j == 1 else WHITE)
            if r_idx == 3:
                set_cell_shading(cells[j], LIGHT_GOLD if j == 2 else (LIGHT_BLUE if j == 1 else WHITE))
            set_cell_text(cells[j], text, color=INK, size=8.7)
    add_para(doc, "Note: Figure is built as an editable Word table so labels and boxes can be adjusted without redrawing the diagram.", size=8.5, color=MUTED, italic=True, after=8)


def add_figure2(doc):
    add_caption(doc, "Figure 2", "AI turn flow")
    table = doc.add_table(rows=1, cols=3)
    apply_table_style(table)
    set_table_geometry(table, [1780, 4080, 3500])
    set_table_borders(table, color="C7D4E2", size="8")
    headers = ["Step", "What AISHA Does", "Main Routes or Outputs"]
    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, BDO_BLUE)
        set_cell_text(cell, h, bold=True, color=WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("1. Receive Turn", "User message or API request is attached to the employee record and simulated date.", "Raw turn saved to SQLite chat memory."),
        ("2. Pre-Route Checks", "AISHA checks whether the message is a pulse reply or a normal chat turn.", "Pulse replies go to scoring. Normal turns go through input guardrails."),
        ("3. Intent Router", "The agent infers what Alyssa is trying to do and chooses tools when needed.", "Knowledge -> search KB\nPlan -> get plan\nTask -> complete task\nPeople -> find person\nSensitive gap -> escalate to HR"),
        ("4. Response + Safety", "The agent composes an answer, then output guardrails enforce citations and redact number-shaped PII.", "Grounded answer, refusal, or escalation response."),
        ("5. Persist + Render", "Streamlit or FastAPI returns the answer and records metadata.", "Updated memory, plan state, sources panel, escalation toast, JSONL log."),
    ]
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for j, text in enumerate(row):
            fill = LIGHT_BLUE if r_idx % 2 == 0 else WHITE
            if j == 0:
                fill = LIGHT_GOLD
            set_cell_shading(cells[j], fill)
            set_cell_text(cells[j], text, bold=(j == 0), color=BDO_DARK_BLUE if j == 0 else INK, size=8.7)
    add_para(doc, "Note: The figure is intentionally simplified. The detailed behavior lives in the prose and source code, while this table keeps the turn flow editable.", size=8.5, color=MUTED, italic=True, after=8)


def parse_md_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        parts = [p.strip() for p in line.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", p or "") for p in parts):
            rows.append(parts)
        i += 1
    return rows, i


def section_anchor(title: str) -> str:
    key = re.sub(r"[^A-Za-z0-9]+", "_", title).strip("_")
    return f"h_{key[:32]}"


def build_body(doc):
    text = SOURCE_MD.read_text(encoding="utf-8")
    start = text.index("## 1. Business Case")
    body = text[start:]
    body = body.replace("## References Used So Far", "## References")
    lines = body.splitlines()
    paragraph_buffer: list[str] = []
    bookmark_id = 10

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            text = " ".join(x.strip() for x in paragraph_buffer).strip()
            if text:
                add_para(doc, text)
            paragraph_buffer = []

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            i += 1
            continue
        if stripped.startswith("```"):
            flush_paragraph()
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1
            continue
        if stripped.startswith("## "):
            flush_paragraph()
            title = stripped[3:].strip()
            if title == "References":
                add_page_break(doc)
            p = doc.add_paragraph(style="Heading 1")
            run = p.add_run(title)
            set_run_font(run, size=16, color=BDO_BLUE, bold=True)
            add_bookmark(p, section_anchor(title), bookmark_id)
            bookmark_id += 1
            i += 1
            continue
        if stripped.startswith("**Figure 1"):
            flush_paragraph()
            add_figure1(doc)
            i += 1
            continue
        if stripped.startswith("**Figure 2"):
            flush_paragraph()
            add_figure2(doc)
            i += 1
            continue
        if stripped.startswith("|"):
            flush_paragraph()
            rows, i = parse_md_table(lines, i)
            add_caption(doc, "Table 1", "Experiment summary")
            add_markdown_table(doc, rows)
            continue
        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph()


def add_appendix(doc):
    add_page_break(doc)
    p = doc.add_paragraph(style="Heading 1")
    title = "Appendix A. Implementation Evidence Map"
    run = p.add_run(title)
    set_run_font(run, size=16, color=BDO_BLUE, bold=True)
    add_bookmark(p, section_anchor(title), 99)
    add_para(
        doc,
        "This appendix keeps the main write-up readable while preserving the technical evidence behind the implementation claims.",
        color=MUTED,
        size=9.5,
        italic=True,
    )
    rows = [
        ["Claim Area", "Evidence in Project", "Why It Matters"],
        ["RAG with citations", "`ingestion.py`, `retriever.py`, Chroma, citation guardrail tests", "Keeps onboarding answers grounded instead of relying on model memory."],
        ["Agent and tool use", "`agent.py`, `tools.py`, fake tool-calling smoke tests", "Shows AISHA can search, read plans, update tasks, route people, and escalate."],
        ["Persistent state", "`state.py`, SQLite tables, memory tests", "Lets Alyssa's progress and chat context survive beyond a single turn."],
        ["Guardrails", "`guardrails.py`, parse and PII redaction tests", "Keeps the assistant scoped to onboarding and safer with sensitive-looking output."],
        ["Pulse support signals", "`pulse.py`, dashboard behavior, pulse tests", "Supports the privacy thesis: HR sees signals and summaries rather than raw transcripts by default."],
        ["API and observability", "`api.py`, `service.py`, `observability.py`, API and logging tests", "Proves the guarded loop can be reused outside Streamlit and inspected without logging message text."],
    ]
    add_markdown_table(doc, rows)


def build_docx(page_map=None):
    doc = Document(TEMPLATE_DOCX)
    set_styles(doc)
    setup_document_settings(doc)
    keep_cover_and_clear_body(doc)
    update_cover(doc)
    add_contents(doc, page_map or {})
    build_body(doc)
    add_appendix(doc)
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--page-map", type=Path)
    args = parser.parse_args()
    page_map = None
    if args.page_map and args.page_map.exists():
        page_map = json.loads(args.page_map.read_text(encoding="utf-8-sig"))
    out = build_docx(page_map=page_map)
    print(out)


if __name__ == "__main__":
    main()
