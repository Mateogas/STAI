"""Build the AISHA Final Capstone DOCX and its two publication diagrams.

The maintainable source remains docs/FINAL_TECHNICAL_WRITEUP.md.  This script
applies the narrative-proposal document preset and converts the deliberately
small Markdown subset used by that source into a polished Word document.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "FINAL_TECHNICAL_WRITEUP.md"
OUTPUT = ROOT / "STAI-Final-Technical-Write-up.docx"
ASSETS = ROOT / "docs" / "assets"
ARCH_IMAGE = ASSETS / "aisha-final-architecture.png"
CERT_IMAGE = ASSETS / "aisha-certificate-flow.png"

NAVY = "17365D"
BLUE = "2E74B5"
TEAL = "0B7A75"
ORANGE = "C55A11"
INK = "202833"
MUTED = "52606D"
PALE = "F4F6F9"
LIGHT_BLUE = "EAF2F8"
LIGHT_TEAL = "E8F5F3"
WHITE = "FFFFFF"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def rounded_box(draw: ImageDraw.ImageDraw, xy, label: str, fill: str, outline: str,
                *, label_size: int = 29, radius: int = 18, text_fill: str = "#17365D") -> None:
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    fnt = font(label_size, True)
    x1, y1, x2, y2 = xy
    max_width = x2 - x1 - 28
    words = label.split()
    lines: list[str] = []
    line = ""
    for word in words:
        trial = f"{line} {word}".strip()
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width:
            line = trial
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    line_height = label_size + 7
    total = len(lines) * line_height
    y = y1 + ((y2 - y1) - total) / 2
    for item in lines:
        box = draw.textbbox((0, 0), item, font=fnt)
        x = x1 + ((x2 - x1) - (box[2] - box[0])) / 2
        draw.text((x, y), item, font=fnt, fill=text_fill)
        y += line_height


def arrow(draw: ImageDraw.ImageDraw, start, end, color="#52606D", width=5, dashed=False) -> None:
    x1, y1 = start
    x2, y2 = end
    if dashed:
        steps = 16
        for i in range(0, steps, 2):
            a = i / steps
            b = min((i + 1) / steps, 1)
            draw.line((x1 + (x2-x1)*a, y1 + (y2-y1)*a,
                       x1 + (x2-x1)*b, y1 + (y2-y1)*b), fill=color, width=width)
    else:
        draw.line((x1, y1, x2, y2), fill=color, width=width)
    import math
    angle = math.atan2(y2-y1, x2-x1)
    length = 18
    spread = 0.55
    points = [
        (x2, y2),
        (x2-length*math.cos(angle-spread), y2-length*math.sin(angle-spread)),
        (x2-length*math.cos(angle+spread), y2-length*math.sin(angle+spread)),
    ]
    draw.polygon(points, fill=color)


def make_architecture_diagram(path: Path) -> None:
    img = Image.new("RGB", (2600, 1700), "#FFFFFF")
    d = ImageDraw.Draw(img)
    title = font(44, True)
    label = font(28, True)
    small = font(23)
    d.text((90, 50), "AISHA final architecture", font=title, fill="#17365D")
    d.text((90, 105), "Models propose; deterministic services authorize, validate, persist, and disclose.",
           font=small, fill="#52606D")

    lanes = [
        ("USER & INTERFACE BOUNDARY", 80, 180, 530, 1430, "#F4F6F9", "#8FA3B8"),
        ("TRUSTED APPLICATION BOUNDARY", 555, 180, 1350, 1430, "#EAF2F8", "#2E74B5"),
        ("MODEL & EVIDENCE BOUNDARIES", 1375, 180, 2050, 1430, "#E8F5F3", "#0B7A75"),
        ("PRIVATE / OPERATIONS BOUNDARIES", 2075, 180, 2520, 1430, "#FFF4E8", "#C55A11"),
    ]
    for name, x1, y1, x2, y2, fill, outline in lanes:
        d.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=fill, outline=outline, width=4)
        lane_font = font(25 if name.startswith("PRIVATE") else 28, True)
        d.text((x1+24, y1+18), name, font=lane_font, fill=outline)

    # User/interface
    rounded_box(d, (135, 300, 475, 430), "Alyssa / Hire", "#FFFFFF", "#8FA3B8")
    rounded_box(d, (135, 500, 475, 630), "HR / support user", "#FFFFFF", "#8FA3B8")
    rounded_box(d, (135, 750, 475, 880), "Streamlit UI", "#FFFFFF", "#2E74B5")
    rounded_box(d, (135, 980, 475, 1110), "Typed REST API", "#FFFFFF", "#2E74B5")
    arrow(d, (305, 430), (305, 750)); arrow(d, (305, 630), (305, 750))
    arrow(d, (305, 880), (305, 980))

    # Trusted core
    core = [
        ((615, 290, 960, 420), "AishaService"),
        ((1000, 290, 1290, 420), "Qwen input guardrail"),
        ((615, 520, 960, 650), "Typed turn planner"),
        ((1000, 520, 1290, 650), "PolicyTurnEngine"),
        ((615, 770, 960, 920), "Schema + evidence validators"),
        ((1000, 770, 1290, 920), "Deterministic fallback"),
        ((615, 1050, 960, 1190), "Consent + sharing controls"),
        ((1000, 1050, 1290, 1190), "SQLite state"),
    ]
    for box, text in core:
        rounded_box(d, box, text, "#FFFFFF", "#2E74B5", label_size=26)
    arrow(d, (475, 815), (615, 355)); arrow(d, (475, 1045), (615, 355))
    arrow(d, (960, 355), (1000, 355)); arrow(d, (1145, 420), (787, 520))
    arrow(d, (960, 585), (1000, 585)); arrow(d, (1145, 650), (1145, 770))
    arrow(d, (1000, 845), (960, 845)); arrow(d, (787, 920), (787, 1050))
    arrow(d, (960, 1120), (1000, 1120))

    # Models/evidence
    rounded_box(d, (1435, 290, 1985, 430), "Llama 3.2 policy agent", "#FFFFFF", "#0B7A75")
    rounded_box(d, (1435, 510, 1985, 650), "Typed read-only tools", "#FFFFFF", "#0B7A75")
    rounded_box(d, (1435, 750, 1985, 890), "Active Handbook + integrity checks", "#FFFFFF", "#0B7A75", label_size=25)
    rounded_box(d, (1435, 980, 1685, 1120), "Chroma", "#FFFFFF", "#0B7A75")
    rounded_box(d, (1735, 980, 1985, 1120), "Nomic embeddings", "#FFFFFF", "#0B7A75", label_size=24)
    arrow(d, (1290, 585), (1435, 360)); arrow(d, (1710, 430), (1710, 510))
    arrow(d, (1710, 650), (1710, 750)); arrow(d, (1710, 890), (1560, 980))
    arrow(d, (1860, 980), (1860, 890))
    arrow(d, (1435, 360), (1290, 845))

    # Private and ops
    rounded_box(d, (2125, 290, 2470, 420), "Medical upload gates", "#FFFFFF", "#C55A11", label_size=25)
    rounded_box(d, (2125, 500, 2470, 630), "Certificate Agent + two typed tools", "#FFFFFF", "#C55A11", label_size=23)
    rounded_box(d, (2125, 710, 2470, 840), "Local PDF / OCR + deterministic rules", "#FFFFFF", "#C55A11", label_size=22)
    rounded_box(d, (2125, 920, 2470, 1050), "Safe result only", "#FFFFFF", "#C55A11", label_size=25)
    rounded_box(d, (2125, 1160, 2470, 1300), "JSONL -> shipper -> relay -> MLflow", "#FFFFFF", "#C55A11", label_size=22)
    arrow(d, (2297, 420), (2297, 500)); arrow(d, (2297, 630), (2297, 710))
    arrow(d, (2297, 840), (2297, 920)); arrow(d, (2125, 985), (1290, 1120))
    arrow(d, (1290, 1180), (2125, 1230), dashed=True)
    d.text((1435, 1215), "closed metadata only", font=small, fill="#C55A11")

    d.line((80, 1490, 2520, 1490), fill="#D6DEE6", width=3)
    d.text((90, 1520), "Privacy rule", font=label, fill="#17365D")
    d.text((275, 1520), "Raw medical content never enters policy chat, HR views, JSONL, or MLflow.",
           font=small, fill="#202833")
    d.text((90, 1575), "Authority rule", font=label, fill="#17365D")
    d.text((305, 1575), "Only validated evidence, explicit consent, and deterministic mutations can change durable state.",
           font=small, fill="#202833")
    img.save(path, dpi=(220, 220))


def make_certificate_diagram(path: Path) -> None:
    img = Image.new("RGB", (2400, 1180), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((80, 45), "Certificate Check: local processing and explicit sharing", font=font(42, True), fill="#17365D")
    boxes = [
        (90, 240, 405, 390, "Notice + acknowledgement", "#EAF2F8", "#2E74B5"),
        (480, 240, 795, 390, "File preflight gates", "#FFF4E8", "#C55A11"),
        (870, 240, 1185, 390, "Bounded Certificate Agent", "#E8F5F3", "#0B7A75"),
        (1260, 240, 1575, 390, "Two typed tools", "#E8F5F3", "#0B7A75"),
        (1650, 240, 1965, 390, "Local OCR + deterministic rules", "#FFF4E8", "#C55A11"),
        (2040, 240, 2310, 390, "Safe result", "#EAF2F8", "#2E74B5"),
    ]
    for x1, y1, x2, y2, text, fill, outline in boxes:
        rounded_box(d, (x1, y1, x2, y2), text, fill, outline, label_size=25)
    for i in range(len(boxes)-1):
        arrow(d, (boxes[i][2], 315), (boxes[i+1][0], 315))

    d.rounded_rectangle((80, 540, 2320, 1010), radius=24, fill="#F7F9FB", outline="#8FA3B8", width=3)
    d.text((120, 585), "What persists", font=font(30, True), fill="#17365D")
    d.text((500, 585), "typed status, reason codes, timing, share state, and privacy-safe action names",
           font=font(26), fill="#202833")
    d.text((120, 685), "What does not persist", font=font(30, True), fill="#17365D")
    d.text((500, 685), "file bytes, OCR text, extracted medical values, filenames, diagnoses, or private reasoning",
           font=font(26), fill="#202833")
    d.text((120, 785), "What HR can see", font=font(30, True), fill="#17365D")
    d.text((500, 785), "only the currently shared safe validation result; revocation removes visibility",
           font=font(26), fill="#202833")
    d.text((120, 885), "What Complete means", font=font(30, True), fill="#17365D")
    d.text((500, 885), "expected synthetic fields were present and structurally consistent—not authenticity or approval",
           font=font(26), fill="#202833")
    img.save(path, dpi=(220, 220))


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    tag = p_pr.find(qn("w:keepNext"))
    if value and tag is None:
        p_pr.append(OxmlElement("w:keepNext"))


def set_header_content(header, text: str) -> None:
    p = header.paragraphs[0]
    p.clear()
    p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(MUTED)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "D6DEE6")
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("AISHA Final Capstone   |   ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_section(section, landscape=False) -> None:
    section.different_first_page_header_footer = False
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    header_text = "AISHA — AI Support for Hires and Associates"
    for header in (section.header, section.even_page_header, section.first_page_header):
        header.is_linked_to_previous = False
        set_header_content(header, header_text)
    for footer in (section.footer, section.even_page_footer, section.first_page_footer):
        footer.is_linked_to_previous = False
        footer_p = footer.paragraphs[0]
        footer_p.clear()
        add_page_field(footer_p)


def style_document(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.22
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 10),
        ("Heading 1", 16, BLUE, 18, 9),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ):
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    styles["Heading 1"].paragraph_format.page_break_before = False

    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(10.5)
        st.paragraph_format.left_indent = Inches(0.375)
        st.paragraph_format.first_line_indent = Inches(-0.194)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.15

    if "Caption" in styles:
        cap = styles["Caption"]
    else:
        cap = styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    cap.font.name = "Calibri"
    cap.font.size = Pt(9)
    cap.font.bold = True
    cap.font.color.rgb = RGBColor.from_string(NAVY)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(6)
    cap.paragraph_format.keep_with_next = True

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    callout.font.name = "Calibri"
    callout.font.size = Pt(10)
    callout.font.color.rgb = RGBColor.from_string(NAVY)
    callout.paragraph_format.left_indent = Inches(0.2)
    callout.paragraph_format.right_indent = Inches(0.2)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(8)


INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[[^\]]+\]\(https?://[^)]+\))")


def add_hyperlink(paragraph, label: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), BLUE); r_pr.append(color)
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single"); r_pr.append(underline)
    run.append(r_pr)
    text = OxmlElement("w:t"); text.text = label; run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text: str) -> None:
    pos = 0
    for match in INLINE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2]); run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1]); run.italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1]); run.font.name = "Consolas"; run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(NAVY)
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_para(doc: Document, text: str, style=None, *, align=None) -> None:
    p = doc.add_paragraph(style=style)
    add_inline(p, text)
    if align is not None:
        p.alignment = align


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"
    # Record the exact usable section width and zero indent in the table XML.
    # Word already lays these tables out to this width; the explicit geometry
    # also makes the artifact portable and auditable outside Word.
    section = doc.sections[-1]
    width_twips = int((section.page_width - section.left_margin - section.right_margin) / 635)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    hdr = table.rows[0]
    for idx, value in enumerate(rows[0]):
        cell = hdr.cells[idx]
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline(p, value)
        set_keep_with_next(p)
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(8 if cols >= 5 else 9)
    set_repeat_table_header(hdr)
    cant_split = OxmlElement("w:cantSplit")
    hdr._tr.get_or_add_trPr().append(cant_split)
    for ridx, source in enumerate(rows[1:], 1):
        row = table.add_row()
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for cidx in range(cols):
            cell = row.cells[cidx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if ridx % 2 == 0:
                set_cell_shading(cell, PALE)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, source[cidx] if cidx < len(source) else "")
            for run in p.runs:
                run.font.size = Pt(7.4 if cols >= 6 else (8 if cols == 5 else 8.8))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    configure_section(section)
    # Keep the title page free of the running header/footer.
    section.different_first_page_header_footer = True
    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("DE LA SALLE UNIVERSITY–MANILA")
    r.font.name = "Calibri"; r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(NAVY)
    add_para(doc, "College of Computer Studies", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Introduction to Agentic AI · STAI100", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "FINAL CAPSTONE", align=WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("AISHA")
    r.font.name = "Calibri"; r.font.size = Pt(34); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(BLUE)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("AI Support for Hires and Associates")
    r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(NAVY)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(14)
    r = p3.add_run("A local-first, evidence-grounded onboarding assistant\nwith consent-controlled policy support and OCR-based certificate checking")
    r.font.size = Pt(12); r.font.italic = True; r.font.color.rgb = RGBColor.from_string(MUTED)

    line = doc.add_paragraph("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.runs[0].font.color.rgb = RGBColor.from_string(TEAL)
    line.paragraph_format.space_before = Pt(18)

    model = doc.add_paragraph()
    model.alignment = WD_ALIGN_PARAGRAPH.CENTER
    model.paragraph_format.space_before = Pt(8)
    add_inline(model, "Local CPU models: **Llama 3.2 · 3.2B** | **Qwen 2.5 · 3.1B / 7.6B** | **Nomic Embed Text · 137M**")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    add_inline(p, "**Submitted by**\nBon Windel Aquino · Jose Miguel Espinosa\nKarl Matthew Dela Cruz · Johann Casio")
    for run in p.runs:
        run.font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    add_inline(p, "**Instructor:** Kristine Kalaw\n**Term:** Term 3, Academic Year 2025–2026\n**Submission date:** 15 August 2026")

    callout = doc.add_paragraph(style="Callout")
    callout.alignment = WD_ALIGN_PARAGRAPH.CENTER
    callout.paragraph_format.space_before = Pt(28)
    add_inline(callout, "**Educational-use notice.** This fictionalized proof of concept is not affiliated with, endorsed by, or representative of BDO Unibank. It uses no real BDO people, documents, systems, or internal data.")
    # Soft emphasis box.
    p_pr = callout._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), LIGHT_BLUE); p_pr.append(shd)
    pbdr = OxmlElement("w:pBdr"); border = OxmlElement("w:bottom")
    border.set(qn("w:val"), "single"); border.set(qn("w:sz"), "10"); border.set(qn("w:color"), BLUE)
    pbdr.append(border); p_pr.append(pbdr)
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    h = doc.add_paragraph("Contents", style="Heading 1")
    h.paragraph_format.page_break_before = False
    items = [
        "Executive summary", "1. Business case and value proposition", "2. Midterm-to-final evolution",
        "3. Review of Related Literature and model selection", "4. Methodology", "5. System architecture",
        "6. Agentic and course-component integration", "7. Experiments and evaluation",
        "8. Results and discussion", "9. Privacy, safety, and governance", "10. Team contributions",
        "11. Retrospective", "12. Conclusion", "References", "Appendices A–G",
    ]
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(6)
        add_inline(p, item)
    p = doc.add_paragraph(style="Callout")
    add_inline(p, "**Status legend used throughout:** Implemented; tested with fakes or deterministic seams; tested live; degraded; or planned/confirmation required.")
    doc.add_page_break()


def add_landscape_table(doc: Document, rows: list[list[str]]) -> None:
    moved_heading = None
    moved_style = None
    if doc.paragraphs and doc.paragraphs[-1].style.name.startswith("Heading"):
        moved_heading = doc.paragraphs[-1].text
        moved_style = doc.paragraphs[-1].style.name
        element = doc.paragraphs[-1]._element
        element.getparent().remove(element)
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(sec, landscape=True)
    if moved_heading:
        p = doc.add_paragraph(moved_heading, style=moved_style)
        p.paragraph_format.page_break_before = False
    if rows and rows[0] and rows[0][0] == "#" and len(rows) > 10:
        # The component matrix is deliberately split at a semantic row
        # boundary.  Word did not reliably repeat the marked header when the
        # single table flowed to a second landscape page, so create an
        # explicit continuation table with its own header.
        # Leave enough room for the explicit page-break paragraph so it does
        # not spill onto an otherwise blank landscape page.
        add_table(doc, rows[:9])
        doc.add_page_break()
        add_table(doc, [rows[0], *rows[9:]])
    else:
        add_table(doc, rows)
    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(sec2, landscape=False)


def convert_markdown(doc: Document, text: str) -> None:
    # Cover content is represented structurally by add_cover(). Start at the
    # first substantive section and retain the explicit fictional-use notice.
    start = text.index("## Fictionalized educational-use notice")
    lines = text[start:].splitlines()
    i = 0
    mermaid_count = 0
    first_heading = True
    while i < len(lines):
        line = lines[i].rstrip()
        if not line or line == "---":
            i += 1
            continue
        if line.startswith("```mermaid"):
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                i += 1
            mermaid_count += 1
            image_path = ARCH_IMAGE if mermaid_count == 1 else CERT_IMAGE
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            width = Inches(6.45) if mermaid_count == 1 else Inches(6.35)
            picture = p.add_run().add_picture(str(image_path), width=width)
            if mermaid_count == 1:
                alt = ("AISHA final architecture showing user and interface, trusted application, "
                       "model and evidence, and private operations boundaries, including Streamlit, "
                       "REST API, guardrail, planner, policy agent, validators, fallback, Active "
                       "Handbook, Chroma, SQLite, certificate processing, and privacy-safe LLMOps.")
            else:
                alt = ("Certificate Check flow from notice and acknowledgement through file preflight, "
                       "bounded Certificate Agent, typed tools, local OCR and deterministic rules, "
                       "to a safe result with explicit sharing and privacy exclusions.")
            picture._inline.docPr.set("title", f"AISHA diagram {mermaid_count}")
            picture._inline.docPr.set("descr", alt)
            i += 1
            continue
        if line.startswith("```"):
            language = line[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i]); i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            r = p.add_run("\n".join(code_lines)); r.font.name = "Consolas"; r.font.size = Pt(8)
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i]); i += 1
            rows = parse_table(table_lines)
            if rows and len(rows[0]) >= 6:
                add_landscape_table(doc, rows)
            else:
                add_table(doc, rows)
            continue
        if line.startswith("#"):
            match = re.match(r"^(#{1,3})\s+(.+)$", line)
            hashes, title = match.groups()
            if title == "Fictionalized educational-use notice":
                p = doc.add_paragraph(style="Callout")
                add_inline(p, "**Fictionalized educational-use notice**")
                p_pr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), LIGHT_BLUE); p_pr.append(shd)
            else:
                # Executive summary is an H1 in the finished document.
                level = 1 if title == "Executive summary" else min(len(hashes), 3)
                p = doc.add_paragraph(title, style=f"Heading {level}")
                if title in {"References", "Appendices"}:
                    p.paragraph_format.page_break_before = True
                if first_heading:
                    p.paragraph_format.page_break_before = False
                    first_heading = False
            i += 1
            continue
        if line.startswith("> "):
            quote = []
            while i < len(lines) and lines[i].startswith("> "):
                quote.append(lines[i][2:]); i += 1
            p = doc.add_paragraph(style="Callout")
            add_inline(p, " ".join(quote))
            p_pr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), PALE); p_pr.append(shd)
            continue
        if re.match(r"^[-*]\s+", line):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.38)
            p.paragraph_format.first_line_indent = Inches(-0.20)
            p.paragraph_format.space_after = Pt(4)
            add_inline(p, "• " + re.sub(r"^[-*]\s+", "", line))
            i += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            number, item = re.match(r"^(\d+)\.\s+(.+)$", line).groups()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.38)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(4)
            add_inline(p, f"{number}.  {item}")
            i += 1
            continue
        if line.startswith("**Figure "):
            p = doc.add_paragraph(style="Caption")
            add_inline(p, line.strip("*"))
            i += 1
            continue

        # Join wrapped Markdown lines into a cohesive paragraph.
        para = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if (not nxt or nxt == "---" or nxt.startswith("#") or nxt.startswith("|")
                    or nxt.startswith("```") or nxt.startswith("> ")
                    or re.match(r"^[-*]\s+", nxt) or re.match(r"^\d+\.\s+", nxt)
                    or nxt.startswith("**Figure ")):
                break
            para.append(nxt); i += 1
        text_value = " ".join(x.strip() for x in para)
        p = doc.add_paragraph()
        add_inline(p, text_value)
        if text_value.startswith("Bauer,") or (doc.paragraphs and any(h.text == "References" for h in doc.paragraphs[-25:])):
            # Reference entries use a hanging indent; appendix text will reset
            # naturally when a new heading appears.
            previous_heading = next((q.text for q in reversed(doc.paragraphs[:-1]) if q.style.name.startswith("Heading")), "")
            if previous_heading == "References":
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.first_line_indent = Inches(-0.3)
                p.paragraph_format.space_after = Pt(8)


def set_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "AISHA: AI Support for Hires and Associates — Final Technical Write-up"
    props.subject = "STAI100 Final Capstone"
    props.author = "Bon Windel Aquino; Jose Miguel Espinosa; Karl Matthew Dela Cruz; Johann Casio"
    props.keywords = "AISHA, STAI, agentic AI, RAG, OCR, privacy, onboarding"
    props.comments = "Fictionalized educational proof of concept; not affiliated with BDO Unibank."


def build() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    make_architecture_diagram(ARCH_IMAGE)
    make_certificate_diagram(CERT_IMAGE)
    doc = Document()
    # Use explicit odd/even headers so Word's PDF exporter cannot silently
    # suppress the running header on alternating pages.  Every configured
    # section receives identical primary/even content; the cover remains the
    # sole blank first-page header.
    doc.settings.odd_and_even_pages_header_footer = True
    style_document(doc)
    set_core_properties(doc)
    add_cover(doc)
    add_contents(doc)
    convert_markdown(doc, SOURCE.read_text(encoding="utf-8"))
    # Ensure the last line is never isolated as a decorative artifact.
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.widow_control = True
        if paragraph.style.name.startswith("Heading"):
            # Explicitly neutralize stray list/section indentation.  Word's
            # exporter otherwise displaced one top-of-page heading into the
            # crop area after pagination.
            paragraph.paragraph_format.left_indent = Inches(0)
            paragraph.paragraph_format.first_line_indent = Inches(0)
    doc.save(OUTPUT)
    print(f"Created {OUTPUT}")
    print(f"Created {ARCH_IMAGE}")
    print(f"Created {CERT_IMAGE}")


if __name__ == "__main__":
    build()
